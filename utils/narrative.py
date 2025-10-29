from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Dict

import pandas as pd
from docx import Document
from docx.shared import Pt
from jinja2 import Environment, FileSystemLoader, select_autoescape

# Use the same normalization as loader, so mapping names match row index reliably
from utils.data_loader import normalize_col


def _safe_get(row: pd.Series, col_name: str, default="N/A"):
    """
    Safely fetch a value from a row by column name (case-insensitive).
    Returns 'N/A' when missing or NaN.
    """
    if not isinstance(col_name, str) or col_name.strip() == "":
        return default
    # exact match
    if col_name in row.index:
        val = row[col_name]
    else:
        # case-insensitive / normalization-aware fallback
        target = normalize_col(col_name)
        hits = [c for c in row.index if normalize_col(c) == target]
        val = row[hits[0]] if hits else default
    if pd.isna(val):
        return default
    return val


def _fmt_number(x):
    """
    Format numbers with thousand separators and 1 decimal if needed.
    Leave non-numeric values unchanged.
    """
    try:
        if x == "N/A":
            return x
        f = float(str(x).replace(",", "").replace(" ", ""))
        if abs(f) >= 1000 and not f.is_integer():
            return f"{f:,.1f}"
        if abs(f) >= 1000:
            return f"{int(round(f)):,.0f}"
        return f"{f:.1f}" if not f.is_integer() else f"{int(f)}"
    except Exception:
        return x


def _fmt_percent(x):
    """
    Format a numeric (possibly with %) as a percent with 1 decimal.
    """
    try:
        if x == "N/A":
            return x
        s = str(x).replace("%", "").replace(",", "")
        f = float(s)
        return f"{f:.1f}"
    except Exception:
        return x


def compose_context(row: pd.Series, mapping: Dict) -> Dict:
    """
    Build a context dict for the Jinja template from the mapping and the selected row.
    """
    ident = mapping.get("identity", {})
    geoid_col = ident.get("geoid_col", "GeoUID")
    name_col = ident.get("name_col", "Geographic name")

    def r(col: str) -> str:
        # exact
        if col in row.index:
            return col
        # normalization-aware match
        target = normalize_col(col)
        for c in row.index:
            if normalize_col(c) == target:
                return c
        return col  # _safe_get will handle "N/A"

    context = {
        "geoid": str(_safe_get(row, r(geoid_col), "N/A")),
        "community_name": str(_safe_get(row, r(name_col), "N/A")),
    }

    def section(key_map):
        out = {}
        for k, col in key_map.items():
            val = _safe_get(row, r(col), "N/A")
            sval = str(val)
            if "%" in k or sval.endswith("%") or "pct" in k or "percentage" in k:
                out[k] = _fmt_percent(val)
            elif any(token in k for token in ["median", "avg", "mean", "count", "total", "number"]):
                out[k] = _fmt_number(val)
            else:
                out[k] = _fmt_number(val)
        return out

    ctx_sections = {
        "pop": section(mapping.get("population_and_dwellings", {})),
        "age": section(mapping.get("age_characteristics", {})),
        "hh": section(mapping.get("household_and_dwelling_characteristics", {})),
        "hht": section(mapping.get("household_type", {})),
        "inc": section(mapping.get("income_of_households_2020", {})),
        "lii": section(mapping.get("low_income_inequality_2020", {})),
        "lang": section(mapping.get("knowledge_of_official_languages", {})),
        "fols": section(mapping.get("first_official_language_spoken", {})),
        "mt": section(mapping.get("mother_tongue", {})),
        "hl": section(mapping.get("language_most_often_at_home", {})),
        "imm": section(mapping.get("immigrant_status_and_period", {})),
        "pbi": section(mapping.get("places_of_birth_immigrants", {})),
        "pr": section(mapping.get("places_of_birth_recent_immigrants", {})),
        "ip": section(mapping.get("indigenous_population", {})),
        "ia": section(mapping.get("indigenous_ancestry", {})),
        "vm": section(mapping.get("visible_minority", {})),
        "ss": section(mapping.get("secondary_school_completion", {})),
        "hc": section(mapping.get("highest_credential", {})),
        "m1": section(mapping.get("mobility_status_1_year", {})),
        "m5": section(mapping.get("mobility_status_5_years", {})),
        "commute": section(mapping.get("main_mode_of_commuting", {})),
        "cd": section(mapping.get("commuting_duration", {})),
        "mol": section(mapping.get("children_eligible_minority_official_language", {})),
        "emol": section(mapping.get("eligibility_instruction_minority_official_language", {})),
    }

    context.update(ctx_sections)
    return context


def render_markdown(context: Dict, lang: str = "en") -> str:
    """
    Render the narrative using Jinja2 markdown template.
    """
    template_file = "narrative_en.md.j2"  # extend later if multi-language
    env = Environment(
        loader=FileSystemLoader(str(Path("templates"))),
        autoescape=select_autoescape(enabled_extensions=("md", "j2")),
        trim_blocks=True,
        lstrip_blocks=True,
    )
    tmpl = env.get_template(template_file)
    return tmpl.render(**context)


def markdown_to_docx(md_text: str) -> bytes:
    """
    Very lightweight Markdown -> DOCX converter good enough for headings and paragraphs.
    - '#', '##', '###' become headings
    - '**bold**' and '*italic*' markers are removed (no inline styles)
    - lines starting with '- ' become bullet paragraphs
    - '---' inserts a page break
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_text.splitlines()

    for line in lines:
        stripped = line.strip("\n")

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(stripped[2:].strip(), style=None)
            p.style = doc.styles["List Paragraph"]
            continue
        if stripped == "---":
            doc.add_page_break()
            continue

        # inline bold/italic (very basic)
        text = stripped
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        doc.add_paragraph(text)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
